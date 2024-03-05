"""
gerando certificados
"""
from docx import Document
from docx.shared import Pt

arquivoWord = Document("Certificado1.docx")

estilo = arquivoWord.styles["Normal"]

for paragrafo in arquivoWord.paragraphs:
    if "@nome" in paragrafo.text:
        paragrafo.text = "Rayssa Ellen da Silva Dantas"
        fonte = estilo.font
        fonte.name = "Calibri (corpo)"
        fonte.size = Pt(24)

arquivoWord.save("Rayssa.docx")


"""
gerando certificados
"""
from docx import Document
from docx.shared import Pt
from openpyxl import load_workbook
import os

nome_arquivo_alunos = "C:\\Users\\Tecnologia\\Documents\\GitHub\\RPA\\Alunos.xlsx"
planilhaDadosAlunos = load_workbook(nome_arquivo_alunos)
sheet_selecionada = planilhaDadosAlunos["Nomes"]

for i in range(2, len(sheet_selecionada["A"]) + 1):
    arquivoWord = Document("Certificado1.docx")

    estilo = arquivoWord.styles["Normal"]

    nomeAluno = sheet_selecionada['A%s' % i].value

    for paragrafo in arquivoWord.paragraphs:
        if "@nome" in paragrafo.text:
            paragrafo.text = nomeAluno
            fonte = estilo.font
            fonte.name = "Calibri (corpo)"
            fonte.size = Pt(24)

    caminhoCertificado = "C:\\Users\\Tecnologia\\Documents\\GitHub\\RPA\\Certificados\\" + nomeAluno + ".docx"
    arquivoWord.save(caminhoCertificado)

print("Certificados gerados com sucesso!")